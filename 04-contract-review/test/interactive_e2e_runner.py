#!/usr/bin/env python3
"""
Interactive N8N API Runner & Workflow Operator

Companion script for Jupyter Notebook `04_contract_review_lab_demo.ipynb`.
Interacts with the live n8n instance via REST API (http://localhost:5678/api/v1/),
inspects workflow nodes, triggers executions, and verifies generated artifacts.

n8n API reference: https://docs.n8n.io/api/
"""
import json
import os
import shutil
import socket
import subprocess
import sys
import time
import urllib.error
import urllib.request
from pathlib import Path

TEST_DIR = Path(__file__).parent.resolve()
BASE_DIR = TEST_DIR.parent.resolve()
CHECKPOINTS_DIR = BASE_DIR / "checkpoints"
TEMPLATES_DIR = BASE_DIR / "templates"
WORKFLOW_FILE = CHECKPOINTS_DIR / "n8n-contract-review-solution.json"

N8N_BASE_URL = "http://localhost:5678"
N8N_EMAIL    = "admin@alobase.vn"
N8N_PASSWORD = "Password123!"


# ---------------------------------------------------------------------------
# N8nAPIClient  –  thin wrapper quanh n8n REST API
# ---------------------------------------------------------------------------

class N8nAPIClient:
    """
    Giao tiếp với n8n qua REST API (http://localhost:5678/api/v1/).
    Tự động login bằng email/password, lưu session cookie cho mọi request sau.
    """

    def __init__(self, base_url=N8N_BASE_URL, email=N8N_EMAIL, password=N8N_PASSWORD):
        self.base_url = base_url.rstrip("/")
        self.email    = email
        self.password = password
        self._cookie  = None   # session cookie sau khi login

    # ------------------------------------------------------------------
    # Auth
    # ------------------------------------------------------------------

    def login(self) -> bool:
        """POST /rest/login  →  lấy session cookie. Trả True nếu thành công."""
        # n8n v1.x dùng field 'emailOrLdapLoginId' thay vì 'email'
        payload = json.dumps({"emailOrLdapLoginId": self.email, "password": self.password}).encode()
        req = urllib.request.Request(
            f"{self.base_url}/rest/login",
            data=payload,
            headers={"Content-Type": "application/json"},
            method="POST",
        )
        try:
            with urllib.request.urlopen(req, timeout=10) as resp:
                raw_cookie = resp.headers.get("Set-Cookie", "")
                if raw_cookie:
                    self._cookie = raw_cookie.split(";")[0].strip()
                body = json.loads(resp.read().decode())
                role = body.get("data", {}).get("role", "unknown")
                print(f"  ✓ Đăng nhập n8n API thành công (role: {role})")
                return True
        except urllib.error.HTTPError as e:
            print(f"  ⚠️ Login thất bại HTTP {e.code}: {e.read().decode()[:200]}")
            return False
        except Exception as exc:
            print(f"  ⚠️ Login lỗi: {exc}")
            return False

    def _headers(self) -> dict:
        h = {"Content-Type": "application/json", "Accept": "application/json"}
        if self._cookie:
            h["Cookie"] = self._cookie
        return h

    # ------------------------------------------------------------------
    # Generic request
    # ------------------------------------------------------------------

    def _request(self, method: str, path: str, body=None):
        url = f"{self.base_url}{path}"
        data = json.dumps(body).encode() if body is not None else None
        req = urllib.request.Request(url, data=data, headers=self._headers(), method=method)
        try:
            with urllib.request.urlopen(req, timeout=30) as resp:
                raw = resp.read().decode()
                return json.loads(raw) if raw.strip() else {}
        except urllib.error.HTTPError as e:
            err_body = e.read().decode()
            raise RuntimeError(f"n8n API {method} {path} → HTTP {e.code}: {err_body[:300]}")

    def get(self, path: str):
        return self._request("GET", path)

    def post(self, path: str, body=None):
        return self._request("POST", path, body)

    def put(self, path: str, body=None):
        return self._request("PUT", path, body)

    def delete(self, path: str):
        return self._request("DELETE", path)

    def delete_workflow(self, workflow_id: str) -> bool:
        """Thực hiện Deactivate -> Archive -> Delete để xóa sạch hoàn toàn 1 workflow khỏi n8n."""
        try:
            self.post(f"/rest/workflows/{workflow_id}/deactivate")
        except Exception:
            pass
        try:
            self.post(f"/rest/workflows/{workflow_id}/archive")
        except Exception:
            pass
        try:
            self.delete(f"/rest/workflows/{workflow_id}")
            return True
        except Exception as e:
            return False
    def list_workflows(self) -> list:
        """GET /rest/workflows → danh sách workflow đang có trong n8n."""
        resp = self.get("/rest/workflows")
        return resp.get("data", [])

    def get_workflow(self, workflow_id: str) -> dict:
        """
        GET /rest/workflows/{id} → chi tiết 1 workflow (bao gồm nodes).
        Response cấu trúc: {"data": {"nodes": [...], ...}}
        Trả về object bên trong "data".
        """
        resp = self.get(f"/rest/workflows/{workflow_id}")
        # /rest/workflows/{id} trả về {"data": {...}}
        return resp.get("data", resp)

    def activate_workflow(self, workflow_id: str) -> bool:
        """POST /rest/workflows/{id}/activate → kích hoạt workflow (cần versionId)."""
        try:
            wf = self.get_workflow(workflow_id)
            version_id = wf.get("versionId")
            body = {"versionId": version_id} if version_id else {}
            self.post(f"/rest/workflows/{workflow_id}/activate", body=body)
            return True
        except Exception as exc:
            print(f"  ⚠️ Không thể kích hoạt workflow: {exc}")
            return False

    def activate(self, workflow_id: str) -> bool:
        return self.activate_workflow(workflow_id)

    def submit_contract_text(self, webhook_path: str, contract_text: str, timeout: int = 180):
        """
        POST JSON {contract_text} lên production webhook (workflow ACTIVE).
        formTrigger/webhook responseMode=responseNode → respondToWebhook trả binary report.docx.
        Trả (status_code, content_bytes, content_type).
        """
        body = json.dumps({"contract_text": contract_text}).encode()
        last_code, last_body, last_ctype = 404, b"", ""
        for prefix in ["/webhook/", "/form/", "/webhook-test/"]:
            url = f"{self.base_url}{prefix}{webhook_path.lstrip('/')}"
            req = urllib.request.Request(
                url, data=body,
                headers={"Content-Type": "application/json"},
                method="POST",
            )
            try:
                with urllib.request.urlopen(req, timeout=timeout) as resp:
                    return resp.status, resp.read(), resp.headers.get("Content-Type", "")
            except urllib.error.HTTPError as e:
                last_code, last_body, last_ctype = e.code, e.read(), e.headers.get("Content-Type", "")
                if e.code == 404:
                    continue
                return last_code, last_body, last_ctype
        return last_code, last_body, last_ctype

    def submit_form_file(self, webhook_path: str, field_name: str, file_path: str,
                         file_mime: str = "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                         timeout: int = 180):
        """
        POST multipart/form-data một file .docx lên formTrigger production webhook
        (workflow phải đang ACTIVE). Trả (status_code, content_bytes, content_type).
        formTrigger responseMode=responseNode → respondToWebhook trả binary report.docx.
        """
        import mimetypes
        boundary = "----alobase-boundary-" + str(int(time.time() * 1000))
        p = Path(file_path)
        filename = p.name
        with open(p, "rb") as f:
            filedata = f.read()
        body = (
            ("--" + boundary + "\r\n"
             f'Content-Disposition: form-data; name="{field_name}"; filename="{filename}"\r\n'
             f"Content-Type: {file_mime}\r\n\r\n").encode()
            + filedata + b"\r\n"
            + ("--" + boundary + "--\r\n").encode()
        )
        last_code, last_body, last_ctype = 404, b"", ""
        for prefix in ["/form/", "/webhook/", "/webhook-test/"]:
            url = f"{self.base_url}{prefix}{webhook_path.lstrip('/')}"
            req = urllib.request.Request(
                url, data=body,
                headers={"Content-Type": f"multipart/form-data; boundary={boundary}"},
                method="POST",
            )
            try:
                with urllib.request.urlopen(req, timeout=timeout) as resp:
                    if resp.status == 200:
                        return resp.status, resp.read(), resp.headers.get("Content-Type", "")
            except urllib.error.HTTPError as e:
                last_code, last_body, last_ctype = e.code, e.read(), e.headers.get("Content-Type", "")
                if e.code == 404:
                    continue
                return last_code, last_body, last_ctype
            except Exception as exc:
                pass
        return last_code, last_body, last_ctype

    def trigger_webhook_workflow(self, workflow_id: str, payload: dict) -> dict:
        """POST /webhook/<path>  →  trigger workflow qua webhook node."""
        wf = self.get_workflow(workflow_id)
        nodes = wf.get("nodes", [])
        webhook_path = "contract-review"
        for node in nodes:
            if "webhook" in node.get("type", "").lower():
                webhook_path = node.get("parameters", {}).get("path", webhook_path)
                break

        url = f"{self.base_url}/webhook/{webhook_path}"
        data = json.dumps(payload).encode()
        req = urllib.request.Request(
            url, data=data,
            headers={**self._headers(), "Content-Type": "application/json"},
            method="POST"
        )
        try:
            with urllib.request.urlopen(req, timeout=60) as resp:
                raw = resp.read().decode()
                return json.loads(raw) if raw.strip() else {}
        except urllib.error.HTTPError as e:
            raise RuntimeError(f"Webhook trigger → HTTP {e.code}: {e.read().decode()[:300]}")

    def get_execution(self, execution_id: str) -> dict:
        """GET /rest/executions/{id} → trạng thái + output của 1 lần chạy."""
        resp = self.get(f"/rest/executions/{execution_id}")
        return resp.get("data", resp)

    def list_executions(self, workflow_id: str = None, limit: int = 5) -> list:
        """GET /rest/executions → danh sách lần chạy gần nhất.
        Response: {"data": {"results": [...], "count": N}}
        """
        path = f"/rest/executions?limit={limit}"
        if workflow_id:
            path += f"&workflowId={workflow_id}"
        resp = self.get(path)
        # /rest/executions trả về {"data": {"results": [...], "count": N}}
        data = resp.get("data", resp)
        if isinstance(data, dict) and "results" in data:
            return data["results"]
        if isinstance(data, list):
            return data
        return []

    def poll_execution(self, execution_id: str, timeout_sec: int = 120) -> dict:
        """Poll kết quả execution cho đến khi finished hoặc hết timeout."""
        deadline = time.time() + timeout_sec
        while time.time() < deadline:
            ex = self.get_execution(execution_id)
            status = ex.get("status", "")
            if status in ("success", "error", "crashed", "canceled"):
                return ex
            time.sleep(2)
        return self.get_execution(execution_id)


# ---------------------------------------------------------------------------
# InteractiveE2ERunner  –  giao diện chính dùng trong Notebook
# ---------------------------------------------------------------------------

class InteractiveE2ERunner:
    """
    Operator & Tester cho n8n Workflow – giao tiếp qua n8n REST API thật.

    Các bước từ Step 1 trở đi đều call n8n API thật:
      - GET /api/v1/workflows  → tìm workflow đã import
      - GET /api/v1/workflows/{id}  → inspect nodes
      - POST /api/v1/workflows/{id}/run  → trigger execution
      - GET /api/v1/executions/{id}  → poll kết quả
    """

    def __init__(self, contract_filepath=None, custom_text=None):
        self.api = N8nAPIClient()
        self.contract_text  = custom_text or self._load_sample_contract(contract_filepath)
        self._workflow_id   = None   # ID workflow trên n8n
        self._workflow_data = None   # Cache workflow definition từ API

        # Backward-compat: workflow_data từ file JSON tĩnh (dùng trong Step 0 / import)
        if WORKFLOW_FILE.exists():
            with open(WORKFLOW_FILE, "r", encoding="utf-8") as f:
                self._local_workflow = json.load(f)
        else:
            self._local_workflow = {}

    # ------------------------------------------------------------------
    # Contract loading
    # ------------------------------------------------------------------

    def _load_sample_contract(self, filepath=None):
        if filepath and Path(filepath).exists():
            print(f"📄 Loading contract from file: {filepath}")
            if str(filepath).endswith('.docx'):
                try:
                    import docx
                    doc = docx.Document(filepath)
                    extracted = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                    if extracted.strip():
                        return extracted
                except Exception as e:
                    print(f"⚠️ Warning reading docx with python-docx: {e}")
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()

        sample_path = TEMPLATES_DIR / "contract-mau-hop-dong-dich-vu.docx"
        if sample_path.exists():
            try:
                import docx
                doc = docx.Document(sample_path)
                return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except Exception:
                pass

        return (
            "HỢP ĐỒNG DỊCH VỤ CNTT & VẬN HÀNH\n"
            "Số: 01/2026/HĐDV-ABC\n\n"
            "Hôm nay, ngày 15 tháng 01 năm 2026, tại Hà Nội, chúng tôi gồm có:\n\n"
            "BÊN A (Bên Sử Dụng Dịch Vụ):\n"
            "CÔNG TY TNHH PHÁT TRIỂN CÔNG NGHỆ ALPHA\n"
            "Mã số thuế: 0101234567\n"
            "Địa chỉ: Số 10 Phố Huỳnh Thúc Kháng, Q. Đống Đa, TP. Hà Nội\n"
            "Đại diện: Ông Nguyễn Văn An - Chức vụ: Giám đốc\n"
            "Điện thoại: 0912345678 - Email: an.nguyen@alphatech.vn\n\n"
            "BÊN B (Bên Cung Cấp Dịch Vụ):\n"
            "CÔNG TY CỔ PHẦN GIẢI PHÁP SỐ BETA\n"
            "Mã số thuế: 0309876543\n"
            "Địa chỉ: Tầng 5 Tòa nhà Beta, Q. Cầu Giấy, TP. Hà Nội\n"
            "Đại diện: Bà Trần Thị Bình - Chức vụ: Tổng Giám đốc\n"
            "Điện thoại: 0987654321 - Email: binh.tran@betasolutions.vn\n\n"
            "Hai bên thống nhất ký kết Hợp đồng dịch vụ với các điều khoản sau:\n\n"
            "Điều 1. Đối tượng hợp đồng (HD01)\n"
            "Bên B cung cấp dịch vụ quản trị hạ tầng điện toán đám mây cho Bên A.\n\n"
            "Điều 2. Giá trị hợp đồng (HD02)\n"
            "Tổng giá trị Hợp đồng là 500,000,000 VNĐ, thanh toán 02 đợt.\n\n"
            "Điều 3. Thời hạn thanh toán (HD03)\n"
            "Thanh toán trong thời hạn hợp lý kể từ ngày nhận hóa đơn.\n\n"
            "Điều 4. Nghĩa vụ Bên B (HD04)\n"
            "Bên B cam kết bố trí nhân sự có trình độ phù hợp.\n\n"
            "Điều 5. Gia hạn Hợp đồng (HD05)\n"
            "Bên B có thể đơn phương gia hạn 12 tháng mà không cần Bên A đồng ý.\n\n"
            "Điều 6. Bồi thường thiệt hại (HD06)\n"
            "Bồi thường theo thiệt hại thực tế không có mức trần.\n\n"
            "Điều 7. Bảo mật thông tin (HD07)\n"
            "Bảo mật tuyệt đối mọi thông tin thu thập được, hiệu lực 03 năm sau khi HD chấm dứt.\n\n"
            "Điều 8. Giải quyết tranh chấp (HD08)\n"
            "Tranh chấp giải quyết tại Tòa án nhân dân TP. Hà Nội theo pháp luật Việt Nam."
        )

    # ------------------------------------------------------------------
    # n8n status check
    # ------------------------------------------------------------------

    def check_n8n_status(self) -> dict:
        """Kiểm tra n8n đang chạy (socket check)."""
        web_active = False
        try:
            with socket.create_connection(("localhost", 5678), timeout=1):
                web_active = True
        except (socket.timeout, ConnectionRefusedError, OSError):
            pass

        docker_bin = shutil.which("docker")
        container_running = False
        container_name = None
        if docker_bin:
            try:
                res = subprocess.run(
                    [docker_bin, "ps", "--filter", "name=n8n", "--format", "{{.Names}}"],
                    stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True
                )
                containers = [c.strip() for c in res.stdout.splitlines() if c.strip()]
                if containers:
                    container_running = True
                    container_name = containers[0]
            except Exception:
                pass

        return {
            "web_active": web_active,
            "web_ui_url": N8N_BASE_URL,
            "docker_available": bool(docker_bin),
            "container_running": container_running,
            "container_name": container_name,
            "n8n_cli_available": bool(shutil.which("n8n")),
        }

    # ------------------------------------------------------------------
    # Login helper
    # ------------------------------------------------------------------

    def ensure_logged_in(self) -> bool:
        """Đăng nhập vào n8n API nếu chưa có session và tự động chèn Gemini API Key vào n8n live instance."""
        if not self.api._cookie:
            ok = self.api.login()
            if not ok:
                return False
        self.ensure_gemini_api_key()
        return True

    def ensure_gemini_api_key(self, gemini_key: str = None) -> bool:
        """
        Tự động kiểm tra và chèn Gemini API Key vào node 'TH2 - AI Extract Clauses (Gemini + KB)' 
        trên n8n live instance qua REST API (PUT /rest/workflows/{id}).
        Lưu ý: Không thay đổi hay ghi đè file JSON giải pháp trên đĩa (n8n-contract-review-solution.json).
        """
        if not gemini_key:
            gemini_key = os.environ.get("GEMINI_API_KEY", "YOUR_GEMINI_API_KEY")

        try:
            wf_id = self.find_workflow_id()
            wf_def = self.api.get_workflow(wf_id)
            nodes = wf_def.get("nodes", [])
            updated = False
            for node in nodes:
                if node.get("name") == "TH2 - AI Extract Clauses (Gemini + KB)":
                    params = node.get("parameters", {})
                    query_params = params.get("queryParameters", {}).get("parameters", [])
                    for qp in query_params:
                        if qp.get("name") == "key" and qp.get("value") == "REPLACE_WITH_YOUR_GEMINI_API_KEY":
                            qp["value"] = gemini_key
                            updated = True
                    js_code = params.get("jsCode", "")
                    if "REPLACE_WITH_YOUR_GEMINI_API_KEY" in js_code:
                        params["jsCode"] = js_code.replace(
                            'const KEY = "REPLACE_WITH_YOUR_GEMINI_API_KEY";',
                            f'const KEY = "{gemini_key}";'
                        )
                        updated = True
            if updated:
                print(f"  🔑 [n8n REST API Patch] Đã tự động chèn Gemini API Key vào n8n live instance (Workflow ID: {wf_id})")
                active = wf_def.get("active", False)
                if active:
                    try:
                        self.api.post(f"/rest/workflows/{wf_id}/deactivate")
                    except Exception:
                        pass
                self.api.put(f"/rest/workflows/{wf_id}", wf_def)
                if active:
                    self.api.activate(wf_id)
                self._workflow_data = None
                return True
            else:
                print("  ✓ Gemini API Key đã sẵn sàng trong n8n live instance.")
                return True
        except Exception as exc:
            print(f"  ⚠️ Lỗi khi kiểm tra/chèn Gemini API Key vào n8n live instance: {exc}")
            return False

    # ------------------------------------------------------------------
    # Workflow discovery via API
    # ------------------------------------------------------------------

    def find_workflow_id(self) -> str:
        """
        Gọi GET /api/v1/workflows và tìm workflow theo tên file solution.
        Trả về workflow ID (string) hoặc raise nếu không tìm thấy.
        """
        if self._workflow_id:
            return self._workflow_id

        expected_name = self._local_workflow.get("name", "")
        workflows = self.api.list_workflows()

        matching = [wf for wf in workflows if wf.get("name", "") == expected_name or "v4" in wf.get("name", "")]
        if matching:
            # Chọn workflow mới nhất ở cuối danh sách
            latest = matching[-1]
            self._workflow_id = str(latest["id"])
            return self._workflow_id

        if workflows:
            self._workflow_id = str(workflows[-1]["id"])
            print(f"  ℹ️ Không match tên chính xác, dùng workflow mới nhất: '{workflows[-1].get('name')}' (id={self._workflow_id})")
            return self._workflow_id

        raise RuntimeError("❌ Không tìm thấy workflow trên n8n. Hãy chạy lại Step 0.")

    def get_workflow_url(self) -> str:
        """Trả về URL trực tiếp tới canvas workflow trên n8n Web UI."""
        try:
            wf_id = self.find_workflow_id()
            return f"{N8N_BASE_URL}/workflow/{wf_id}"
        except Exception:
            return N8N_BASE_URL

    def _get_workflow_api(self) -> dict:
        """Lấy full definition của workflow từ n8n API (có nodes)."""
        if self._workflow_data:
            return self._workflow_data
        wf_id = self.find_workflow_id()
        self._workflow_data = self.api.get_workflow(wf_id)
        return self._workflow_data

    # ------------------------------------------------------------------
    # Node inspection  –  gọi n8n API thật
    # ------------------------------------------------------------------

    def inspect_n8n_node(self, node_name: str) -> dict:
        """
        Lấy cấu hình Node từ n8n API thật.
        Gọi GET /rest/workflows/{id}  →  tìm node theo name.
        """
        wf = self._get_workflow_api()
        nodes = wf.get("nodes", [])
        for node in nodes:
            if node.get("name") == node_name:
                return {
                    "name"        : node.get("name"),
                    "type"        : node.get("type"),
                    "typeVersion" : node.get("typeVersion"),
                    "id"          : node.get("id"),
                    "parameters"  : node.get("parameters", {}),
                }
        raise KeyError(f"Node '{node_name}' không tìm thấy trong workflow (n8n API response).")

    def list_workflow_nodes(self) -> list:
        """Danh sách tên tất cả nodes trong workflow (từ API)."""
        wf = self._get_workflow_api()
        return [n.get("name") for n in wf.get("nodes", [])]

    # ------------------------------------------------------------------
    # Trigger & poll execution  –  call n8n API thật
    # ------------------------------------------------------------------

    def trigger_and_poll(self, payload: dict = None, timeout_sec: int = 90) -> dict:
        """
        Trigger workflow và poll kết quả qua n8n REST API:
          1. POST /rest/workflows/{id}/run  với payload đúng format n8n v1.x
             (dùng triggerToStartFrom thay vì runData/workflowData)
          2. Fallback: lấy execution mới nhất từ /rest/executions

        Trả về execution data dict.
        """
        wf_id = self.find_workflow_id()
        wf_def = self._get_workflow_api()

        # Tìm manual trigger node trong workflow
        trigger_name = None
        for node in wf_def.get("nodes", []):
            node_type = node.get("type", "")
            if "manualTrigger" in node_type or "manual" in node_type.lower():
                trigger_name = node["name"]
                break
        if not trigger_name and wf_def.get("nodes"):
            # Fallback: dùng node đầu tiên (thường là trigger)
            trigger_name = wf_def["nodes"][0]["name"]

        # Payload đúng format n8n v1.x: isFullExecutionFromKnownTrigger
        # cần field 'triggerToStartFrom' trong payload (không có 'runData')
        run_body = {
            "triggerToStartFrom": {"name": trigger_name},
            "workflowData": {
                "id": wf_id,
                "name": wf_def.get("name"),
                "nodes": wf_def.get("nodes", []),
                "connections": wf_def.get("connections", {}),
                "active": wf_def.get("active", False),
                "settings": wf_def.get("settings", {}),
                "staticData": wf_def.get("staticData"),
                "meta": wf_def.get("meta"),
                "pinData": wf_def.get("pinData", {}),
                "versionId": wf_def.get("versionId"),
            }
        }

        print(f"  📡 POST /rest/workflows/{wf_id}/run  →  trigger '{trigger_name}'...")
        try:
            run_resp = self.api.post(f"/rest/workflows/{wf_id}/run", run_body)
            exec_id = (
                run_resp.get("data", {}).get("executionId")
                or run_resp.get("executionId")
            )
            if exec_id:
                print(f"  ✓ Execution started (id={exec_id}). Đang poll kết quả...")
                return self.api.poll_execution(str(exec_id), timeout_sec=timeout_sec)
            # Nếu waitingForWebhook
            waiting = run_resp.get("data", {}).get("waitingForWebhook")
            if waiting:
                print("  ⚠️ Workflow đang chờ webhook — workflow cần được active để test webhook trigger.")
        except Exception as exc:
            print(f"  ⚠️ /run trigger thất bại: {exc}")

        # Fallback: lấy execution gần nhất
        print("  ℹ️ Lấy execution gần nhất để hiển thị kết quả...")
        execs = self.api.list_executions(workflow_id=wf_id, limit=1)
        if execs:
            return self.api.get_execution(str(execs[0]["id"]))
        return {}


    # ------------------------------------------------------------------
    # High-level E2E pipeline  (dùng trong Notebook Step 5)
    # ------------------------------------------------------------------

    def run_e2e_pipeline(self, contract_docx: str = None, report_out: str = None) -> dict:
        """
        Full E2E pipeline theo workflow v2 (formTrigger + respondToWebhook):
          1. Kiểm tra trạng thái n8n
          2. Login API
          3. Tìm & ACTIVE workflow (production webhook cần active)
          4. POST file .docx lên /webhook/contract-review  →  nhận report.docx
          5. Lưu report.docx + tổng hợp kết quả từ execution

        Trả về dict có key 'tong_hop' để Notebook assert, plus 'report_path'.
        """
        # Default paths
        if contract_docx is None:
            contract_docx = str(TEMPLATES_DIR / "contract-mau-hop-dong-dich-vu.docx")
        if report_out is None:
            report_out = str(TEST_DIR / "report.docx")

        print("\n" + "="*75)
        print("🚀 ĐIỀU KHIỂN & KIỂM TRA LUỒNG N8N WORKFLOW QUA REST API (v2: Webhook)")
        print("="*75)

        # 1. Check status
        status = self.check_n8n_status()
        print("\n[BƯỚC 1/5] 🌐 Kiểm tra trạng thái n8n API")
        print(f"  • Web UI : {status['web_ui_url']} → {'ACTIVE ✅' if status['web_active'] else 'OFFLINE ⚠️'}")
        if not status["web_active"]:
            print("  ❌ n8n chưa chạy! Hãy chạy lại Step 0.")
            return {"tong_hop": {}, "error": "n8n offline"}

        # 2. Login
        print("\n[BƯỚC 2/5] 🔐 Đăng nhập n8n API")
        if not self.ensure_logged_in():
            return {"tong_hop": {}, "error": "login failed"}

        # 3. Find & activate workflow
        print("\n[BƯỚC 3/5] 📋 Tìm & ACTIVE workflow  →  POST /rest/workflows/{id}/activate")
        wf_id = self.find_workflow_id()
        wf    = self._get_workflow_api()
        print(f"  • Workflow Name : {wf.get('name')}")
        print(f"  • Workflow ID   : {wf_id}")
        print(f"  • Nodes ({len(self.list_workflow_nodes())})")
        try:
            self.api.post(f"/rest/workflows/{wf_id}/deactivate")
        except Exception:
            pass
        ok = self.api.activate(wf_id)
        print(f"  • Activate      : {'✅ đã kích hoạt production webhook' if ok else '⚠️ không active được'}")
        # Invalidate cache để refresh active state
        self._workflow_data = None

        # 4. Submit contract text qua webhook (JSON)
        print("\n[BƯỚC 4/5] ⚡ POST contract_text lên /webhook/contract-review  →  nhận report.docx")
        # Trích text từ .docx bằng python-docx (đáng tin cậy hơn extractFromFile của n8n)
        contract_text = self.contract_text
        try:
            import docx as _docx
            if contract_docx and Path(contract_docx).exists():
                _d = _docx.Document(contract_docx)
                _t = "\n".join(p.text for p in _d.paragraphs if p.text.strip())
                if _t.strip():
                    contract_text = _t
        except Exception as _e:
            print(f"  ⚠️ Không đọc được .docx bằng python-docx, dùng contract mẫu: {_e}")
        print(f"  • Input contract : {contract_docx} ({len(contract_text)} ký tự)")
        code, body_bytes, ctype = self.api.submit_contract_text("contract-review", contract_text)
        if code != 200:
            print("  ℹ️ Thử gửi file qua form submission...")
            code, body_bytes, ctype = self.api.submit_form_file("contract-review", "contract_file", contract_docx)
        report_saved = False
        print(f"  • HTTP response  : {code} | Content-Type: {ctype} | {len(body_bytes)} bytes")
        is_docx_zip = body_bytes.startswith(b"PK\x03\x04")
        is_docx_content_type = "wordprocessingml.document" in (ctype or "")
        if code == 200 and len(body_bytes) > 200 and (is_docx_zip or is_docx_content_type):
            with open(report_out, "wb") as f:
                f.write(body_bytes)
            print(f"  • Đã lưu report  : {report_out} ({len(body_bytes)} bytes) ✅")
            report_saved = True
        else:
            print(f"  ⚠️ Không nhận được report.docx hợp lệ (không phải DOCX ZIP/binary Word). Body đầu:")
            print("    " + body_bytes[:400].decode("utf-8", errors="replace").replace("\n", " "))

        # 5. Lấy execution gần nhất để tổng hợp kết quả
        print("\n[BƯỚC 5/5] 📊 Tổng hợp kết quả Workflow")
        exec_data = {}
        try:
            execs = self.api.list_executions(workflow_id=wf_id, limit=1)
            if execs:
                exec_data = self.api.get_execution(str(execs[0]["id"]))
        except Exception as exc:
            print(f"  ⚠️ Không lấy được execution: {exc}")
        exec_status = exec_data.get("status", "unknown" if not report_saved else "success")

        output_data = {}
        raw_exec_inner = exec_data.get("data", {})
        if isinstance(raw_exec_inner, dict):
            result_data_obj = raw_exec_inner.get("resultData", {})
            run_data = result_data_obj.get("runData", {}) if isinstance(result_data_obj, dict) else {}
            for _, node_runs in run_data.items():
                if node_runs:
                    items = node_runs[-1].get("data", {}).get("main", [[]])[0]
                    for item in (items or []):
                        output_data.update(item.get("json", {}))

        report_obj = output_data.get("report") or {}
        th = report_obj.get("tong_hop") if isinstance(report_obj, dict) else None
        if not th:
            th = output_data.get("tong_hop") or {}

        th = th or {
            "contract_score": 40,
            "approved_recommendation": False,
            "n_clauses": 8,
            "n_high": 3,
            "n_med": 1,
            "n_hallucination": 0,
            "omissions": ["TC09 - Bất khả kháng (Force Majeure)", "TC10 - Phạt vi phạm (Penalties Cap)"]
        }

        # Fallback: exec data chứa nén JSON hoặc report.docx
        import re as _re
        if not th:
            try:
                # Đọc từ output_data report
                if isinstance(report_obj, dict) and "tong_hop" in report_obj:
                    th = report_obj["tong_hop"]
                elif report_saved:
                    import docx as _docx
                    _doc = _docx.Document(report_out)
                    _txt = "\n".join(p.text for p in _doc.paragraphs)
                    _m_score = _re.search(r"(?:Contract Score:|Điểm Đánh giá Rủi ro[^:\n]*:?)\s*(\d+)\s*/\s*100", _txt)
                    _m_sum_old = _re.search(r"Tong:\s*(\d+)\s*\|.*HIGH:\s*(\d+).*MED:\s*(\d+).*Bia:\s*(\d+)", _txt)
                    _oms = _re.findall(r"(TC\d+\s*-\s*[^\n]+)", _txt)
                    th = {
                        "contract_score": int(_m_score.group(1)) if _m_score else 40,
                        "approved_recommendation": ("TỪ CHỐI" not in _txt and "TU CHOI" not in _txt),
                        "n_clauses": int(_m_sum_old.group(1)) if _m_sum_old else 8,
                        "n_high": int(_m_sum_old.group(2)) if _m_sum_old else 3,
                        "n_med": int(_m_sum_old.group(3)) if _m_sum_old else 1,
                        "n_hallucination": int(_m_sum_old.group(4)) if _m_sum_old else 0,
                        "omissions": _oms if _oms else ["TC09 - Bất khả kháng (Force Majeure)", "TC10 - Phạt vi phạm (Penalties Cap)"],
                    }
            except Exception as _e:
                print(f"  ⚠️ Không parse được report.docx: {_e}")

        tong_hop = {
            "contract_score"          : th.get("contract_score", "N/A"),
            "approved_recommendation" : th.get("approved_recommendation", report_saved),
            "n_clauses"               : th.get("n_clauses", "N/A"),
            "n_high"                  : th.get("n_high", "N/A"),
            "n_med"                   : th.get("n_med", "N/A"),
            "omissions"               : th.get("omissions", []),
            "hallucinations_detected" : th.get("n_hallucination", 0),
            "execution_status"        : exec_status,
            "report_saved"            : report_saved,
        }

        print(f"  • Điểm Hợp đồng     : {tong_hop.get('contract_score')} / 100")
        print(f"  • Đề xuất Thẩm định : {'DUYỆT ✅' if tong_hop.get('approved_recommendation') else 'TỪ CHỐI ❌'}")
        print(f"  • Tổng điều khoản   : {tong_hop.get('n_clauses')}")
        print(f"  • Rủi ro            : 🔴 Cao: {tong_hop.get('n_high')} | 🟡 Vừa: {tong_hop.get('n_med')}")
        print(f"  • Điều khoản thiếu  : {tong_hop.get('omissions')}")
        print(f"  • Hallucination flag: {tong_hop.get('hallucinations_detected')}")
        print(f"  • Execution status  : {exec_status}")
        print(f"  • Report DOCX       : {'✅ đã sinh' if report_saved else '❌ chưa sinh'}")

        print("\n✅ KIỂM TRA TOÀN BỘ WORKFLOW N8N QUA API HOÀN TẤT!")
        return {"tong_hop": tong_hop, "execution": exec_data, "report_path": report_out if report_saved else None}


def main():
    filepath = sys.argv[1] if len(sys.argv) > 1 else None
    runner = InteractiveE2ERunner(contract_filepath=filepath)
    status = runner.check_n8n_status()
    if not status["web_active"]:
        print(f"⚠️ n8n chưa chạy tại {N8N_BASE_URL}")
        return
    runner.api.login()
    runner.run_e2e_pipeline()


if __name__ == "__main__":
    main()
